"""Convert a Markdown manuscript to a .docx Word document using pypandoc."""

from __future__ import annotations

import logging
import os
import re
import subprocess
import tempfile
import zipfile
from pathlib import Path

import pypandoc
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _convert_svg_to_png(svg_path: Path, png_path: Path) -> None:
    """Convert one SVG to PNG using cairo/rsvg tools."""
    png_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        import cairosvg  # type: ignore[import-not-found]

        cairosvg.svg2png(url=str(svg_path), write_to=str(png_path))
        return
    except Exception:
        pass

    try:
        subprocess.run(
            ["rsvg-convert", str(svg_path), "-o", str(png_path)],
            check=True,
            capture_output=True,
            text=True,
        )
        return
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"conversion from SVG failed: {exc}") from exc


def _rewrite_svg_image_links_for_docx(markdown: str, source_dir: Path, work_dir: Path) -> str:
    """Rewrite markdown image links ending in .svg to generated PNG fallbacks."""

    image_pattern = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
    converted: dict[str, str] = {}

    def _replace(match: re.Match[str]) -> str:
        alt_text = match.group(1)
        target = match.group(2).strip()
        path_token = target.split()[0].strip()
        if not path_token.lower().endswith(".svg"):
            return match.group(0)
        if path_token in converted:
            return f"![{alt_text}]({converted[path_token]})"

        svg_path = Path(path_token)
        if not svg_path.is_absolute():
            svg_path = (source_dir / svg_path).resolve()
        if not svg_path.exists():
            logger.warning("DOCX export: SVG image missing, leaving as-is: %s", path_token)
            return match.group(0)

        png_name = f"{Path(path_token).stem}_docx.png"
        png_path = work_dir / png_name
        try:
            _convert_svg_to_png(svg_path, png_path)
        except Exception as exc:  # noqa: BLE001
            logger.warning('Could not convert image %s: "%s"', path_token, exc)
            return f"*Figure omitted in DOCX export (SVG conversion failed): {Path(path_token).name}*"
        converted[path_token] = png_name
        return f"![{alt_text}]({png_name})"

    return image_pattern.sub(_replace, markdown)


def _run_pandoc_to_docx(source_path: Path, output_path: Path, resource_paths: list[Path]) -> None:
    """Run pandoc conversion with a joined resource-path list."""
    joined_resource_path = os.pathsep.join(str(p) for p in resource_paths)
    pypandoc.convert_file(
        str(source_path),
        "docx",
        outputfile=str(output_path),
        extra_args=[
            "--standalone",
            "--resource-path",
            joined_resource_path,
            "--dpi=300",  # preserve image resolution in DOCX (pandoc default is 96)
        ],
    )


def _fix_table_layout(docx_path: Path) -> None:
    """Post-process all tables in *docx_path* for correct rendering in Word and Google Docs.

    Steps applied to every table:
    1. Apply the built-in 'Table Grid' style for visible borders (ECMA-376 compliant,
       survives DOCX->Google Docs conversion without extra XML border manipulation).
    2. Force 100% page width via w:tblW type="pct" w="5000".
    3. Force fixed layout via w:tblLayout type="fixed" so Word respects explicit widths.
    4. Rescale w:gridCol widths so they sum to TEXT_WIDTH_TWIPS (6.5" text area),
       preserving the proportions pandoc derived from the markdown separator dashes.
    5. Sync each cell's w:tcW to its column's scaled width (cells override gridCol
       when tcW is present, so both must match).
    """
    TEXT_WIDTH_TWIPS = 9360  # 6.5" * 1440 twips/inch (Letter page, 1" margins)

    _BORDER_SIDES = ("top", "bottom", "start", "end", "insideH", "insideV")

    doc = Document(str(docx_path))
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # Step 1: add visible borders via w:tblBorders (ECMA-376 OOXML, survives Google Docs)
        tblBorders = tblPr.find(qn("w:tblBorders"))
        if tblBorders is None:
            tblBorders = OxmlElement("w:tblBorders")
            tblPr.append(tblBorders)
        for side in _BORDER_SIDES:
            el = tblBorders.find(qn(f"w:{side}"))
            if el is None:
                el = OxmlElement(f"w:{side}")
                tblBorders.append(el)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")  # 6/8 = 0.75 pt
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")

        # Step 2: 100% page width
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")

        # Step 3: fixed layout
        tblLayout = tblPr.find(qn("w:tblLayout"))
        if tblLayout is None:
            tblLayout = OxmlElement("w:tblLayout")
            tblPr.append(tblLayout)
        tblLayout.set(qn("w:type"), "fixed")

        # Step 4: rescale gridCol widths proportionally to fill text area
        grid_cols = tbl.tblGrid.findall(qn("w:gridCol"))
        scaled: list[int] = []
        if grid_cols:
            total = sum(int(g.get(qn("w:w"), "0") or "0") for g in grid_cols)
            if total > 0:
                for col in grid_cols:
                    raw = int(col.get(qn("w:w"), "0") or "0")
                    new_w = int(round(raw / total * TEXT_WIDTH_TWIPS))
                    col.set(qn("w:w"), str(new_w))
                    scaled.append(new_w)

        # Step 5: sync each cell's tcW to its column's scaled width
        if scaled:
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx >= len(scaled):
                        break
                    tc = cell._tc
                    tcPr = tc.find(qn("w:tcPr"))
                    if tcPr is None:
                        tcPr = OxmlElement("w:tcPr")
                        tc.insert(0, tcPr)
                    tcW = tcPr.find(qn("w:tcW"))
                    if tcW is None:
                        tcW = OxmlElement("w:tcW")
                        tcPr.append(tcW)
                    tcW.set(qn("w:w"), str(scaled[idx]))
                    tcW.set(qn("w:type"), "dxa")

        # Step 6: explicitly bold the header row so it is visually distinct in all
        # renderers (Word, Google Docs, Cursor). Pandoc's embedded "Table" style
        # only adds a bottom border to the first row -- no bold, no background --
        # making the header visually identical to data rows without this step.
        if table.rows:
            for cell in table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        rPr = run._r.find(qn("w:rPr"))
                        if rPr is None:
                            rPr = OxmlElement("w:rPr")
                            run._r.insert(0, rPr)
                        for tag in ("w:b", "w:bCs"):
                            if rPr.find(qn(tag)) is None:
                                el = OxmlElement(tag)
                                rPr.append(el)

    doc.save(str(docx_path))


def generate_docx(source_path: Path, output_path: Path) -> Path:
    """Convert the Markdown file at *source_path* to a Word .docx at *output_path*.

    Uses pypandoc.convert_file() so pandoc resolves relative image paths
    (e.g. fig_prisma_flow.png) against the source file's directory.
    This is the key difference from convert_text(): text strings have no
    directory anchor, so relative image references silently fail.

    Returns *output_path* on success; raises RuntimeError on failure.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    abs_source = source_path.resolve()
    source_markdown = abs_source.read_text(encoding="utf-8")
    with tempfile.TemporaryDirectory(prefix="docx_export_") as temp_dir_str:
        temp_dir = Path(temp_dir_str)
        rewritten_markdown = _rewrite_svg_image_links_for_docx(source_markdown, abs_source.parent, temp_dir)
        if rewritten_markdown != source_markdown:
            staged_markdown = temp_dir / "docx_input.md"
            staged_markdown.write_text(rewritten_markdown, encoding="utf-8")
            _run_pandoc_to_docx(staged_markdown, output_path, [temp_dir, abs_source.parent])
        else:
            _run_pandoc_to_docx(abs_source, output_path, [abs_source.parent])

    if not output_path.exists():
        raise RuntimeError(f"pypandoc did not produce output at {output_path}")
    if not zipfile.is_zipfile(output_path):
        raise RuntimeError(f"Pandoc produced a non-DOCX artifact at {output_path}")
    _fix_table_layout(output_path)
    return output_path
