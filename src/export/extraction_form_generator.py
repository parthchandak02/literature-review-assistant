"""
Data Extraction Form Generator

Generates data extraction form templates for systematic reviews.
"""

import json
from pathlib import Path
from typing import Dict, List, Any
import logging

from ..schemas.extraction_schemas import ExtractedDataSchema

logger = logging.getLogger(__name__)


class ExtractionFormGenerator:
    """Generates data extraction form templates."""

    def __init__(self):
        """Initialize extraction form generator."""
        self.fields = self._get_extraction_fields()

    def _get_extraction_fields(self) -> List[Dict[str, Any]]:
        """Get list of extraction fields from schema."""
        schema = ExtractedDataSchema.model_json_schema()
        properties = schema.get("properties", {})

        fields = []
        for field_name, field_info in properties.items():
            field_type = field_info.get("type", "string")
            description = field_info.get("description", "")

            fields.append(
                {
                    "name": field_name,
                    "type": field_type,
                    "description": description,
                    "required": field_name in schema.get("required", []),
                }
            )

        return fields

    def generate_form(
        self,
        output_path: str,
        format: str = "markdown",
    ) -> str:
        """
        Generate data extraction form.

        Args:
            output_path: Path to save form file
            format: Output format ("markdown", "json", "word")

        Returns:
            Path to generated form file
        """
        output_path_obj = Path(output_path)
        output_path_obj.parent.mkdir(parents=True, exist_ok=True)

        if format == "markdown":
            return self._generate_markdown_form(output_path_obj)
        elif format == "json":
            return self._generate_json_form(output_path_obj)
        elif format == "word":
            return self._generate_word_form(output_path_obj)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _generate_markdown_form(self, output_path: Path) -> str:
        """Generate markdown form."""
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("# Data Extraction Form\n\n")
            f.write("Use this form to extract data from included studies.\n\n")
            f.write("---\n\n")

            # Basic metadata section
            f.write("## Basic Metadata\n\n")
            basic_fields = ["title", "authors", "year", "journal", "doi"]
            for field_name in basic_fields:
                field = next((f for f in self.fields if f["name"] == field_name), None)
                if field:
                    f.write(f"### {field_name.replace('_', ' ').title()}\n\n")
                    f.write(f"**Description:** {field['description']}\n\n")
                    f.write("**Value:** ___________________________\n\n")

            # Study characteristics section
            f.write("## Study Characteristics\n\n")
            study_fields = [
                "study_objectives",
                "methodology",
                "study_design",
                "participants",
                "interventions",
                "outcomes",
                "country",
                "setting",
                "sample_size",
            ]
            for field_name in study_fields:
                field = next((f for f in self.fields if f["name"] == field_name), None)
                if field:
                    f.write(f"### {field_name.replace('_', ' ').title()}\n\n")
                    f.write(f"**Description:** {field['description']}\n\n")
                    if field["type"] == "array":
                        f.write("**Value:**\n")
                        f.write("- [ ] Item 1\n")
                        f.write("- [ ] Item 2\n\n")
                    elif field["type"] == "integer":
                        f.write("**Value:** ___________________________\n\n")
                    else:
                        f.write("**Value:**\n\n")
                        f.write("___________________________________________\n\n")

            # Results section
            f.write("## Results\n\n")
            results_fields = ["key_findings", "detailed_outcomes", "quantitative_results"]
            for field_name in results_fields:
                field = next((f for f in self.fields if f["name"] == field_name), None)
                if field:
                    f.write(f"### {field_name.replace('_', ' ').title()}\n\n")
                    f.write(f"**Description:** {field['description']}\n\n")
                    f.write("**Value:**\n\n")
                    f.write("___________________________________________\n\n")

            # Limitations
            f.write("## Limitations\n\n")
            f.write("### limitations\n\n")
            f.write("**Description:** Study limitations\n\n")
            f.write("**Value:**\n\n")
            f.write("___________________________________________\n\n")

        logger.info(f"Markdown extraction form generated at {output_path}")
        return str(output_path)

    def _generate_json_form(self, output_path: Path) -> str:
        """Generate JSON form template."""
        form_template = {
            "form_name": "Data Extraction Form",
            "version": "1.0",
            "fields": [],
        }

        for field in self.fields:
            form_template["fields"].append(
                {
                    "name": field["name"],
                    "type": field["type"],
                    "description": field["description"],
                    "required": field["required"],
                    "value": None,
                }
            )

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(form_template, f, indent=2)

        logger.info(f"JSON extraction form generated at {output_path}")
        return str(output_path)

    def _generate_word_form(self, output_path: Path) -> str:
        """Generate Word form (requires python-docx)."""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not available, falling back to markdown")
            return self._generate_markdown_form(output_path.with_suffix(".md"))

        doc = Document()
        doc.add_heading("Data Extraction Form", level=0)
        doc.add_paragraph("Use this form to extract data from included studies.")

        # Add fields
        for field in self.fields:
            doc.add_heading(field["name"].replace("_", " ").title(), level=2)
            doc.add_paragraph(field["description"])
            doc.add_paragraph("Value: ___________________________")
            doc.add_paragraph()

        doc.save(str(output_path))
        logger.info(f"Word extraction form generated at {output_path}")
        return str(output_path)
