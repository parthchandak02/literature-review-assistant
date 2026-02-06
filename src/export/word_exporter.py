"""
Word Exporter

Exports systematic review reports to Microsoft Word format (.docx) for journal submission.
"""

from pathlib import Path
from typing import Dict, List, Optional, Any
import logging

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)


class WordExporter:
    """Exports reports to Word format."""

    def __init__(self):
        """Initialize Word exporter."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")

    def export(
        self,
        report_data: Dict[str, Any],
        output_path: str,
        journal: Optional[str] = None,
    ) -> str:
        """
        Export report to Word format.

        Args:
            report_data: Dictionary containing report sections and metadata
            output_path: Path to save Word file
            journal: Journal name (for formatting)

        Returns:
            Path to generated Word file
        """
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        # Create document
        doc = Document()

        # Title
        title = report_data.get("title", "Systematic Review")
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Author information
        authors = report_data.get("authors", [])
        if authors:
            author_names = [a.get("name", "") for a in authors]
            author_para = doc.add_paragraph(", ".join(author_names))
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Abstract
        abstract = report_data.get("abstract", "")
        if abstract:
            doc.add_heading("Abstract", level=1)
            doc.add_paragraph(abstract)

        # Keywords
        keywords = report_data.get("keywords", [])
        if keywords:
            doc.add_heading("Keywords", level=1)
            keywords_text = ", ".join(keywords)
            doc.add_paragraph(keywords_text)

        # Introduction
        introduction = report_data.get("introduction", "")
        if introduction:
            doc.add_heading("Introduction", level=1)
            self._add_markdown_text(doc, introduction)

        # Methods
        methods = report_data.get("methods", "")
        if methods:
            doc.add_heading("Methods", level=1)
            self._add_markdown_text(doc, methods)

        # Results
        results = report_data.get("results", "")
        if results:
            doc.add_heading("Results", level=1)
            self._add_markdown_text(doc, results)

        # Discussion
        discussion = report_data.get("discussion", "")
        if discussion:
            doc.add_heading("Discussion", level=1)
            self._add_markdown_text(doc, discussion)

        # Figures
        figures = report_data.get("figures", [])
        for figure in figures:
            fig_path = figure.get("path", "")
            caption = figure.get("caption", "Figure")
            if fig_path and Path(fig_path).exists():
                doc.add_picture(fig_path, width=Inches(6))
                para = doc.add_paragraph(caption)
                para.style = "Caption"

        # References
        references = report_data.get("references", [])
        if references:
            doc.add_heading("References", level=1)
            for ref in references:
                ref_text = self._format_reference(ref)
                doc.add_paragraph(ref_text, style="List Paragraph")

        # Save document
        doc.save(str(output_file))
        logger.info(f"Word file generated: {output_file}")
        return str(output_file)

    def _add_markdown_text(self, doc: Document, text: str):
        """Add markdown-formatted text to document."""
        if not text:
            return

        # Simple markdown parsing
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                doc.add_paragraph()
                i += 1
                continue

            # Check for markdown table
            if line.startswith("|") and "---" in lines[i+1] if i+1 < len(lines) else False:
                # Parse table
                table = self._parse_markdown_table(lines, i)
                if table:
                    self._add_table(doc, table)
                    # Skip table lines
                    i += len(table["rows"]) + 2  # Header + separator + rows
                    continue

            # Headers
            if line.startswith("###"):
                doc.add_heading(line[3:].strip(), level=3)
            elif line.startswith("##"):
                doc.add_heading(line[2:].strip(), level=2)
            elif line.startswith("#"):
                doc.add_heading(line[1:].strip(), level=1)
            else:
                # Regular paragraph
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)
            
            i += 1

    def _parse_markdown_table(self, lines: List[str], start_idx: int) -> Optional[Dict[str, Any]]:
        """Parse markdown table from lines."""
        if start_idx >= len(lines):
            return None

        header_line = lines[start_idx]
        if not header_line.startswith("|"):
            return None

        # Parse header
        headers = [cell.strip() for cell in header_line.split("|")[1:-1]]

        # Check separator
        if start_idx + 1 >= len(lines):
            return None
        separator_line = lines[start_idx + 1]
        if "---" not in separator_line:
            return None

        # Parse rows
        rows = []
        i = start_idx + 2
        while i < len(lines):
            row_line = lines[i]
            if not row_line.startswith("|"):
                break
            row_cells = [cell.strip() for cell in row_line.split("|")[1:-1]]
            if len(row_cells) == len(headers):
                rows.append(row_cells)
            i += 1

        return {"headers": headers, "rows": rows}

    def _add_table(self, doc: Document, table: Dict[str, Any]):
        """Add table to Word document."""

        num_cols = len(table["headers"])
        num_rows = len(table["rows"])

        word_table = doc.add_table(rows=num_rows + 1, cols=num_cols)
        word_table.style = "Light Grid Accent 1"

        # Add header row
        header_cells = word_table.rows[0].cells
        for i, header in enumerate(table["headers"]):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Add data rows
        for row_idx, row_data in enumerate(table["rows"]):
            cells = word_table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(cells):
                    cells[col_idx].text = cell_data

    def _add_formatted_text(self, para, text: str):
        """Add formatted text to paragraph."""
        # Handle bold **text**
        import re
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)

    def _format_reference(self, ref: Dict[str, Any]) -> str:
        """Format a single reference."""
        authors = ref.get("authors", [])
        title = ref.get("title", "")
        journal = ref.get("journal", "")
        year = ref.get("year", "")
        doi = ref.get("doi", "")

        parts = []
        if authors:
            parts.append(", ".join(authors[:3]))
            if len(authors) > 3:
                parts[-1] += " et al."
        if title:
            parts.append(f'"{title}"')
        if journal:
            parts.append(journal)
        if year:
            parts.append(str(year))
        if doi:
            parts.append(f"doi: {doi}")

        return ". ".join(parts) + "."
