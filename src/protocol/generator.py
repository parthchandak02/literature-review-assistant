"""PROSPERO-format protocol generator."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from src.models import ProsperoRunData, ProtocolDocument, ReviewConfig, SettingsConfig

# Mapping from review_type enum value to a PROSPERO-appropriate study design description.
# PROSPERO item 21 asks for the types of *primary* studies to be included, not the review type itself.
_STUDY_DESIGN_DESCRIPTIONS: dict[str, str] = {
    "systematic": (
        "Non-randomized studies, cohort studies, cross-sectional studies, "
        "quasi-experimental designs, observational studies, and usability evaluations"
    ),
    "scoping": (
        "Any study design including systematic reviews, randomized and non-randomized "
        "trials, observational studies, surveys, and grey literature"
    ),
    "meta_analysis": (
        "Randomized controlled trials and non-randomized studies with quantitative "
        "outcome data suitable for statistical pooling"
    ),
    "narrative": ("Any study design including experimental, observational, qualitative, and mixed-methods studies"),
}


def _format_run_date(run_id: str) -> str:
    """Normalize run_id/timestamp strings to YYYY-MM-DD when possible."""
    if not run_id:
        return "Not provided"
    if re.match(r"^\d{8}(-\d{6})?$", run_id):
        return f"{run_id[0:4]}-{run_id[4:6]}-{run_id[6:8]}"
    if re.match(r"^\d{4}-\d{2}-\d{2}", run_id):
        return run_id[:10]
    return run_id


_PLACEHOLDER_PATTERNS: tuple[str, ...] = (
    r"\[[A-Z_ ]+\]",
    r"keyword one",
    r"describe the",
    r"what is the effect of",
    r"\[Not specified\]",
)


_NOT_PROVIDED = "Not provided"


def _contains_placeholder_text(value: str) -> bool:
    text = value.strip().lower()
    if not text:
        return True
    return any(re.search(pattern, text, flags=re.IGNORECASE) for pattern in _PLACEHOLDER_PATTERNS)


def _collect_placeholder_warnings(config: ReviewConfig) -> list[str]:
    warnings: list[str] = []
    checks: list[tuple[str, str]] = [
        ("research_question", config.research_question),
        ("domain", config.domain),
        ("pico.population", config.pico.population),
        ("pico.intervention", config.pico.intervention),
        ("pico.comparison", config.pico.comparison),
        ("pico.outcome", config.pico.outcome),
        ("scope", config.scope),
    ]
    for field_name, value in checks:
        if _contains_placeholder_text(value):
            warnings.append(field_name)
    if not config.keywords or any(_contains_placeholder_text(k) for k in config.keywords):
        warnings.append("keywords")
    return warnings


class ProtocolGenerator:
    def __init__(self, output_dir: str = "runs"):
        self.output_dir = Path(output_dir)

    def generate(
        self,
        workflow_id: str,
        config: ReviewConfig,
        settings: SettingsConfig | None = None,
    ) -> ProtocolDocument:
        screening_cfg = getattr(settings, "screening", None)
        batch_threshold = float(getattr(screening_cfg, "batch_screen_threshold", 0.20))
        validation_fraction = float(getattr(screening_cfg, "batch_screen_validation_fraction", 0.10))
        validation_pct = int(round(validation_fraction * 100))
        return ProtocolDocument(
            workflow_id=workflow_id,
            research_question=config.research_question,
            pico=config.pico,
            eligibility_criteria=config.inclusion_criteria + config.exclusion_criteria,
            planned_databases=config.target_databases,
            planned_screening_method=(
                "Three-stage: (1) BM25 keyword pre-filter; "
                f"(2) batch LLM pre-ranker (relevance threshold {batch_threshold:.2f}) "
                f"with {validation_pct}% cross-validation; "
                "(3) independent dual-reviewer screening with adjudication. "
                "Full-text retrieval via multi-tier resolver (Unpaywall, Semantic Scholar, "
                "Europe PMC, CORE, PubMed Central) followed by full-text eligibility assessment."
            ),
            planned_rob_tools=["rob2", "robins_i", "casp", "mmat"],
            planned_synthesis_method=(
                "Meta-analysis when feasible (>=2 studies, homogeneous numeric outcomes); "
                "otherwise narrative synthesis following SWiM 2021 guidelines "
                "(Campbell et al., Syst Rev 2020) with grouping by outcome domain "
                "and direction-of-effect vote-counting."
            ),
            prospero_id=config.protocol.registration_number or None,
        )

    def _render_other_methods_text(self, config: ReviewConfig, run_data: ProsperoRunData | None = None) -> str:
        observed = [m for m in (getattr(run_data, "other_methods_searched", []) or []) if m]
        if observed:
            return "Non-database sources used in this run: " + ", ".join(sorted(set(observed))) + "."
        configured_non_db = [
            src for src in (config.target_databases or []) if src in {"clinicaltrials_gov", "perplexity_search"}
        ]
        if configured_non_db:
            return "Configured non-database sources: " + ", ".join(configured_non_db) + "."
        return "No additional non-database sources were configured for this run."

    def validate_prospero_inputs(self, config: ReviewConfig) -> list[str]:
        """Return fields that still look like placeholders for PROSPERO output."""
        return _collect_placeholder_warnings(config)

    def render_markdown(self, protocol: ProtocolDocument, config: ReviewConfig) -> str:
        sections: list[tuple[str, str]] = [
            ("1. Review title", config.research_question),
            ("2. Original language title", config.research_question),
            ("3. Anticipated start date", _NOT_PROVIDED),
            ("4. Anticipated completion date", _NOT_PROVIDED),
            ("5. Stage of review at time of registration", "Started"),
            ("6. Named contact", _NOT_PROVIDED),
            ("7. Named contact email", _NOT_PROVIDED),
            ("8. Named contact address", _NOT_PROVIDED),
            ("9. Named contact phone", _NOT_PROVIDED),
            ("10. Organisational affiliation", _NOT_PROVIDED),
            ("11. Review team members and affiliations", _NOT_PROVIDED),
            ("12. Funding sources/sponsors", config.funding.source),
            ("13. Conflicts of interest", config.conflicts_of_interest),
            ("14. Review question", protocol.research_question),
            ("15. Searches", ", ".join(protocol.planned_databases)),
            ("16. URL to search strategy", "doc_search_strategies_appendix.md"),
            ("17. Condition/domain being studied", config.domain),
            ("18. Participants/population", config.pico.population),
            ("19. Intervention(s), exposure(s)", config.pico.intervention),
            ("20. Comparator(s)/control", config.pico.comparison),
            (
                "21. Types of study to be included",
                _STUDY_DESIGN_DESCRIPTIONS.get(
                    config.review_type.value,
                    f"Primary studies appropriate for a {config.review_type.value} review",
                ),
            ),
            ("22. Context", config.scope),
        ]
        lines = ["# PROSPERO Protocol Draft", ""]
        for header, body in sections:
            lines.append(f"## {header}")
            lines.append(body)
            lines.append("")
        return "\n".join(lines).strip() + "\n"

    def render_prospero_docx(
        self,
        protocol: ProtocolDocument,
        config: ReviewConfig,
        run_data: ProsperoRunData,
    ) -> Path:
        """Generate a fully-populated PROSPERO registration form as a .docx file.

        The output mirrors the official PROSPERO form section structure.
        Fields that require post-registration assignment (e.g. registration
        number) are left as '[TO BE ASSIGNED]' placeholders.
        """
        doc = Document()

        # --- helpers ---------------------------------------------------------

        def _heading1(text: str) -> None:
            p = doc.add_heading(text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.color.rgb = RGBColor(0, 0, 0)

        def _heading2(text: str) -> None:
            p = doc.add_heading(text, level=2)
            run = p.runs[0] if p.runs else p.add_run(text)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        def _field(label: str, value: str) -> None:
            """Write a labelled field: bold label on first line, value below."""
            p = doc.add_paragraph()
            bold_run = p.add_run(label + ": ")
            bold_run.bold = True
            p.add_run(value if value else _NOT_PROVIDED)

        def _divider() -> None:
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)

        # build derived values once
        db_list = ", ".join(config.target_databases) if config.target_databases else _NOT_PROVIDED
        keywords_str = "; ".join(config.keywords) if config.keywords else _NOT_PROVIDED
        study_design_str = _STUDY_DESIGN_DESCRIPTIONS.get(
            config.review_type.value,
            f"Primary studies appropriate for a {config.review_type.value} review",
        )
        rob_tools_str = ", ".join(t.upper() for t in protocol.planned_rob_tools)
        author_str = config.author_name or _NOT_PROVIDED
        funding_str = config.funding.source if config.funding and config.funding.source else "No external funding"
        coi_str = config.conflicts_of_interest if config.conflicts_of_interest else "None declared"
        outcome_str = config.pico.outcome if config.pico else _NOT_PROVIDED
        scope_str = config.scope or _NOT_PROVIDED
        domain_str = config.domain or _NOT_PROVIDED
        run_date = _format_run_date(run_data.run_id)
        _criteria_blob = " ".join([*config.inclusion_criteria, *config.exclusion_criteria]).lower()
        language_restrictions = "No language restrictions applied."
        if "english" in _criteria_blob:
            language_restrictions = "English-language restriction applied."
        synthesis_str = run_data.synthesis_method or protocol.planned_synthesis_method

        # --- Document title --------------------------------------------------
        title_para = doc.add_heading("PROSPERO Registration Form", level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "International Prospective Register of Systematic Reviews"
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # =====================================================================
        # SECTION 1: REVIEW TITLE AND BASIC DETAILS
        # =====================================================================
        _heading1("REVIEW TITLE AND BASIC DETAILS")
        _divider()

        _heading2("Review title")
        doc.add_paragraph(config.research_question)

        _heading2("Condition or domain being studied")
        doc.add_paragraph(domain_str)

        _heading2("Rationale for the review")
        doc.add_paragraph(
            f"This systematic review addresses the following research question: "
            f"{config.research_question}. "
            f"The review targets the population of {config.pico.population if config.pico else _NOT_PROVIDED} "
            f"and examines the intervention/exposure of "
            f"{config.pico.intervention if config.pico else _NOT_PROVIDED}."
        )

        _heading2("Review objectives")
        doc.add_paragraph(
            f"To systematically identify, appraise, and synthesize evidence addressing: {config.research_question}"
        )

        _heading2("Keywords")
        doc.add_paragraph(keywords_str)

        _heading2("Country")
        doc.add_paragraph(_NOT_PROVIDED)

        _heading2("PROSPERO registration number")
        doc.add_paragraph(_NOT_PROVIDED)

        # =====================================================================
        # SECTION 2: ELIGIBILITY CRITERIA
        # =====================================================================
        doc.add_paragraph()
        _heading1("ELIGIBILITY CRITERIA")
        _divider()

        _heading2("Population")
        _field("Included", config.pico.population if config.pico else _NOT_PROVIDED)

        _heading2("Intervention(s) or exposure(s)")
        _field("Included", config.pico.intervention if config.pico else _NOT_PROVIDED)

        _heading2("Comparator(s) or control(s)")
        _field("Included", config.pico.comparison if config.pico else _NOT_PROVIDED)

        _heading2("Main outcome(s)")
        _field("Included", outcome_str)

        _heading2("Study design")
        doc.add_paragraph(study_design_str)

        _heading2("Inclusion criteria")
        for criterion in config.inclusion_criteria:
            doc.add_paragraph(criterion, style="List Bullet")
        if not config.inclusion_criteria:
            doc.add_paragraph(_NOT_PROVIDED)

        _heading2("Exclusion criteria")
        for criterion in config.exclusion_criteria:
            doc.add_paragraph(criterion, style="List Bullet")
        if not config.exclusion_criteria:
            doc.add_paragraph(_NOT_PROVIDED)

        _heading2("Context")
        doc.add_paragraph(scope_str)

        # =====================================================================
        # SECTION 3: SIMILAR REVIEWS
        # =====================================================================
        doc.add_paragraph()
        _heading1("SIMILAR REVIEWS")
        _divider()

        _heading2("Check for similar records already in PROSPERO")
        doc.add_paragraph(
            "A search of PROSPERO should be conducted prior to registration submission "
            "to identify any similar ongoing or completed reviews. Any identified records "
            "and the rationale for proceeding with this review should be documented here."
        )

        # =====================================================================
        # SECTION 4: TIMELINE OF THE REVIEW
        # =====================================================================
        doc.add_paragraph()
        _heading1("TIMELINE OF THE REVIEW")
        _divider()

        _heading2("Date of first submission to PROSPERO")
        doc.add_paragraph(_NOT_PROVIDED)

        _heading2("Review timeline")
        date_start = str(config.date_range_start) if config.date_range_start else _NOT_PROVIDED
        date_end = str(config.date_range_end) if config.date_range_end else _NOT_PROVIDED
        doc.add_paragraph(f"Start date: {date_start}          End date: {date_end}")

        _heading2("Date of registration in PROSPERO")
        doc.add_paragraph(_NOT_PROVIDED)

        # =====================================================================
        # SECTION 5: AVAILABILITY OF FULL PROTOCOL
        # =====================================================================
        doc.add_paragraph()
        _heading1("AVAILABILITY OF FULL PROTOCOL")
        _divider()

        _heading2("Availability of full protocol")
        doc.add_paragraph(
            "The full systematic review protocol is available as doc_protocol.md in the "
            "run directory. This includes complete search strategies, eligibility criteria, "
            "data extraction forms, and quality assessment procedures."
        )

        # =====================================================================
        # SECTION 6: SEARCHING AND SCREENING
        # =====================================================================
        doc.add_paragraph()
        _heading1("SEARCHING AND SCREENING")
        _divider()

        _heading2("Search for unpublished studies")
        doc.add_paragraph(self._render_other_methods_text(config, run_data))

        _heading2("Main bibliographic databases that will be searched")
        doc.add_paragraph(db_list)
        if run_data.search_counts:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run("Records retrieved per database:").bold = True
            for db in config.target_databases:
                count = run_data.search_counts.get(db, 0)
                doc.add_paragraph(f"  {db}: {count} records", style="List Bullet")

        _heading2("Search language restrictions")
        doc.add_paragraph(language_restrictions)

        _heading2("Search date restrictions")
        doc.add_paragraph(f"Literature published between {date_start} and {date_end}.")

        _heading2("Other methods of identifying studies")
        doc.add_paragraph(
            "Supplementary methods are limited to non-database sources configured for this run "
            "(for example trial registries or grey-literature endpoints)."
        )

        _heading2("Link to search strategy")
        doc.add_paragraph("Full search strategies available in doc_search_strategies_appendix.md in the run directory.")

        _heading2("Selection process")
        doc.add_paragraph(protocol.planned_screening_method)

        _heading2("Other relevant information about searching and screening")
        doc.add_paragraph(
            f"Total records identified: {sum(run_data.search_counts.values()) if run_data.search_counts else '[Not yet available]'}. "
            f"Records after deduplication and screening: {run_data.included_count} studies included. "
            f"Full texts retrieved: {run_data.fulltext_retrieved_count}."
        )

        # =====================================================================
        # SECTION 7: DATA COLLECTION PROCESS
        # =====================================================================
        doc.add_paragraph()
        _heading1("DATA COLLECTION PROCESS")
        _divider()

        _heading2("Data extraction from published articles and reports")
        doc.add_paragraph(
            "Structured data extraction performed independently using an AI-assisted extraction "
            "pipeline with human oversight. Extracted fields include: study design, sample size, "
            "population characteristics, intervention details, comparator details, outcome "
            "measures and results, follow-up duration, and country of study."
        )

        _heading2("Study risk of bias or quality assessment")
        doc.add_paragraph(
            f"Risk of bias assessed using the following validated tools: {rob_tools_str}. "
            "Tool selection based on study design: RoB 2 for randomized trials, "
            "ROBINS-I for non-randomized interventional studies, CASP for observational designs."
        )

        _heading2("Reporting bias assessment")
        doc.add_paragraph(
            "Funnel plot asymmetry assessed where meta-analysis is feasible (>=10 studies). "
            "Egger's test applied for statistical assessment of small-study effects."
        )

        _heading2("Certainty assessment")
        doc.add_paragraph(
            "Evidence certainty assessed using the GRADE (Grading of Recommendations "
            "Assessment, Development and Evaluation) framework for each primary outcome. "
            "Summary of Findings table generated per GRADE guidelines."
        )

        # =====================================================================
        # SECTION 8: OUTCOMES TO BE ANALYSED
        # =====================================================================
        doc.add_paragraph()
        _heading1("OUTCOMES TO BE ANALYSED")
        _divider()

        _heading2("Main outcomes")
        doc.add_paragraph(outcome_str)

        _heading2("Additional outcomes")
        doc.add_paragraph(
            "Secondary outcomes identified during data extraction; "
            "reported narratively where meta-analysis is not feasible."
        )

        # =====================================================================
        # SECTION 9: PLANNED DATA SYNTHESIS
        # =====================================================================
        doc.add_paragraph()
        _heading1("PLANNED DATA SYNTHESIS")
        _divider()

        _heading2("Strategy for data synthesis")
        doc.add_paragraph(synthesis_str)

        # =====================================================================
        # SECTION 10: CURRENT REVIEW STAGE
        # =====================================================================
        doc.add_paragraph()
        _heading1("CURRENT REVIEW STAGE")
        _divider()

        _heading2("Stage of the review at this submission")
        stages = [
            ("Preliminary searches", "YES"),
            ("Piloting of the study selection process", "YES"),
            ("Formal screening of search results against eligibility criteria", "YES"),
            ("Data extraction", "YES"),
            ("Risk of bias (quality) assessment", "YES"),
            ("Data analysis", "YES"),
        ]
        for stage_label, done in stages:
            p = doc.add_paragraph()
            p.add_run(f"[{'X' if done == 'YES' else ' '}] ").bold = True
            p.add_run(stage_label)

        _heading2("Publication of review results")
        doc.add_paragraph(_NOT_PROVIDED)

        # =====================================================================
        # SECTION 11: REVIEW AFFILIATION, FUNDING AND PEER REVIEW
        # =====================================================================
        doc.add_paragraph()
        _heading1("REVIEW AFFILIATION, FUNDING AND PEER REVIEW")
        _divider()

        _heading2("Review team members")
        doc.add_paragraph(author_str)

        _heading2("Review affiliation")
        doc.add_paragraph(_NOT_PROVIDED)

        _heading2("Funding source")
        doc.add_paragraph(funding_str)

        _heading2("Peer review")
        doc.add_paragraph(_NOT_PROVIDED)

        # =====================================================================
        # SECTION 12: ADDITIONAL INFORMATION
        # =====================================================================
        doc.add_paragraph()
        _heading1("ADDITIONAL INFORMATION")
        _divider()

        _heading2("Additional information")
        doc.add_paragraph(
            f"This systematic review was conducted using an automated pipeline "
            f"(Literature Review Assistant). Search conducted on: {run_date}. "
            f"Full run artifacts are available in the run directory."
        )

        _heading2("Review conflict of interest")
        doc.add_paragraph(coi_str)

        _heading2("Medical Subject Headings")
        doc.add_paragraph(keywords_str)

        # --- write file ------------------------------------------------------
        out_path = self.output_dir / "doc_prospero_registration.docx"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return out_path

    def render_prospero_markdown(
        self,
        protocol: ProtocolDocument,
        config: ReviewConfig,
        run_data: ProsperoRunData,
    ) -> str:
        """Render the PROSPERO registration form as markdown."""
        db_list = ", ".join(config.target_databases) if config.target_databases else _NOT_PROVIDED
        keywords_str = "; ".join(config.keywords) if config.keywords else _NOT_PROVIDED
        study_design_str = _STUDY_DESIGN_DESCRIPTIONS.get(
            config.review_type.value,
            f"Primary studies appropriate for a {config.review_type.value} review",
        )
        rob_tools_str = ", ".join(t.upper() for t in protocol.planned_rob_tools)
        author_str = config.author_name or _NOT_PROVIDED
        funding_str = config.funding.source if config.funding and config.funding.source else "No external funding"
        coi_str = config.conflicts_of_interest if config.conflicts_of_interest else "None declared"
        outcome_str = config.pico.outcome if config.pico else _NOT_PROVIDED
        scope_str = config.scope or _NOT_PROVIDED
        domain_str = config.domain or _NOT_PROVIDED
        run_date = _format_run_date(run_data.run_id)
        date_start = str(config.date_range_start) if config.date_range_start else _NOT_PROVIDED
        date_end = str(config.date_range_end) if config.date_range_end else _NOT_PROVIDED
        synthesis_str = run_data.synthesis_method or protocol.planned_synthesis_method
        criteria_blob = " ".join([*config.inclusion_criteria, *config.exclusion_criteria]).lower()
        language_restrictions = "No language restrictions applied."
        if "english" in criteria_blob:
            language_restrictions = "English-language restriction applied."
        placeholder_warnings = _collect_placeholder_warnings(config)

        lines: list[str] = [
            "# PROSPERO Registration Form",
            "",
            "International Prospective Register of Systematic Reviews",
            "",
            "## REVIEW TITLE AND BASIC DETAILS",
            "",
            "### Review title",
            config.research_question,
            "",
            "### Condition or domain being studied",
            domain_str,
            "",
            "### Rationale for the review",
            (
                f"This systematic review addresses the following research question: {config.research_question}. "
                f"The review targets the population of {config.pico.population if config.pico else _NOT_PROVIDED} "
                f"and examines the intervention/exposure of "
                f"{config.pico.intervention if config.pico else _NOT_PROVIDED}. "
            ),
            "",
            "### Review objectives",
            f"To systematically identify, appraise, and synthesize evidence addressing: {config.research_question}",
            "",
            "### Keywords",
            keywords_str,
            "",
            "### Country",
            _NOT_PROVIDED,
            "",
            "### PROSPERO registration number",
            _NOT_PROVIDED,
            "",
            "## ELIGIBILITY CRITERIA",
            "",
            "### Population",
            f"- Included: {config.pico.population if config.pico else _NOT_PROVIDED}",
            "",
            "### Intervention(s) or exposure(s)",
            f"- Included: {config.pico.intervention if config.pico else _NOT_PROVIDED}",
            "",
            "### Comparator(s) or control(s)",
            f"- Included: {config.pico.comparison if config.pico else _NOT_PROVIDED}",
            "",
            "### Main outcome(s)",
            f"- Included: {outcome_str}",
            "",
            "### Study design",
            study_design_str,
            "",
            "### Inclusion criteria",
        ]
        if config.inclusion_criteria:
            lines.extend([f"- {c}" for c in config.inclusion_criteria])
        else:
            lines.append(_NOT_PROVIDED)
        lines.extend(
            [
                "",
                "### Exclusion criteria",
            ]
        )
        if config.exclusion_criteria:
            lines.extend([f"- {c}" for c in config.exclusion_criteria])
        else:
            lines.append(_NOT_PROVIDED)
        lines.extend(
            [
                "",
                "### Context",
                scope_str,
                "",
                "## SIMILAR REVIEWS",
                "",
                "### Check for similar records already in PROSPERO",
                (
                    "A search of PROSPERO should be conducted prior to registration submission "
                    "to identify any similar ongoing or completed reviews. Any identified records "
                    "and the rationale for proceeding with this review should be documented here."
                ),
                "",
                "## TIMELINE OF THE REVIEW",
                "",
                "### Date of first submission to PROSPERO",
                _NOT_PROVIDED,
                "",
                "### Review timeline",
                f"Start date: {date_start}          End date: {date_end}",
                "",
                "### Date of registration in PROSPERO",
                _NOT_PROVIDED,
                "",
                "## AVAILABILITY OF FULL PROTOCOL",
                "",
                "### Availability of full protocol",
                (
                    "The full systematic review protocol is available as doc_protocol.md in the run directory. "
                    "This includes complete search strategies, eligibility criteria, data extraction forms, "
                    "and quality assessment procedures."
                ),
                "",
                "## SEARCHING AND SCREENING",
                "",
                "### Search for unpublished studies",
                self._render_other_methods_text(config, run_data),
                "",
                "### Main bibliographic databases that will be searched",
                db_list,
                "",
            ]
        )
        if run_data.search_counts:
            lines.extend(
                [
                    "#### Records retrieved per database",
                    *[f"- {db}: {run_data.search_counts.get(db, 0)} records" for db in config.target_databases],
                    "",
                ]
            )
        lines.extend(
            [
                "### Search language restrictions",
                language_restrictions,
                "",
                "### Search date restrictions",
                f"Literature published between {date_start} and {date_end}.",
                "",
                "### Other methods of identifying studies",
                (
                    "Supplementary methods are limited to non-database sources configured for this run "
                    "(for example trial registries or grey-literature endpoints)."
                ),
                "",
                "### Link to search strategy",
                "Full search strategies available in doc_search_strategies_appendix.md in the run directory.",
                "",
                "### Selection process",
                protocol.planned_screening_method,
                "",
                "### Other relevant information about searching and screening",
                (
                    f"Total records identified: {sum(run_data.search_counts.values()) if run_data.search_counts else '[Not yet available]'}. "
                    f"Records after deduplication and screening: {run_data.included_count} studies included. "
                    f"Full texts retrieved: {run_data.fulltext_retrieved_count}."
                ),
                "",
                "## DATA COLLECTION PROCESS",
                "",
                "### Data extraction from published articles and reports",
                (
                    "Structured data extraction performed independently using an AI-assisted extraction pipeline "
                    "with human oversight. Extracted fields include: study design, sample size, population "
                    "characteristics, intervention details, comparator details, outcome measures and results, "
                    "follow-up duration, and country of study."
                ),
                "",
                "### Study risk of bias or quality assessment",
                (
                    f"Risk of bias assessed using the following validated tools: {rob_tools_str}. "
                    "Tool selection based on study design: RoB 2 for randomized trials, ROBINS-I for "
                    "non-randomized interventional studies, CASP for observational designs."
                ),
                "",
                "### Reporting bias assessment",
                (
                    "Reporting-bias checks are applied where quantitative synthesis is feasible and sufficient "
                    "studies are available."
                ),
                "",
                "### Certainty assessment",
                (
                    "Evidence certainty assessed using the GRADE (Grading of Recommendations Assessment, "
                    "Development and Evaluation) framework for each primary outcome. Summary of Findings "
                    "table generated per GRADE guidelines."
                ),
                "",
                "## OUTCOMES TO BE ANALYSED",
                "",
                "### Main outcomes",
                outcome_str,
                "",
                "### Additional outcomes",
                "Secondary outcomes identified during data extraction; reported narratively when quantitative pooling is not feasible.",
                "",
                "## PLANNED DATA SYNTHESIS",
                "",
                "### Strategy for data synthesis",
                synthesis_str,
                "",
                "## CURRENT REVIEW STAGE",
                "",
                "### Stage of the review at this submission",
                "- [x] Preliminary searches",
                "- [x] Piloting of the study selection process",
                "- [x] Formal screening of search results against eligibility criteria",
                "- [x] Data extraction",
                "- [x] Risk of bias (quality) assessment",
                "- [x] Data analysis",
                "",
                "### Publication of review results",
                _NOT_PROVIDED,
                "",
                "## REVIEW AFFILIATION, FUNDING AND PEER REVIEW",
                "",
                "### Review team members",
                author_str,
                "",
                "### Review affiliation",
                _NOT_PROVIDED,
                "",
                "### Funding source",
                funding_str,
                "",
                "### Peer review",
                _NOT_PROVIDED,
                "",
                "## ADDITIONAL INFORMATION",
                "",
                "### Additional information",
                (
                    f"This systematic review was conducted using an automated pipeline (Literature Review Assistant). "
                    f"Search conducted on: {run_date}. Full run artifacts are available in the run directory."
                ),
                "",
                "### Review conflict of interest",
                coi_str,
                "",
                "### Medical Subject Headings",
                keywords_str,
                "",
            ]
        )
        if placeholder_warnings:
            lines.extend(
                [
                    "## PRE-SUBMISSION CHECKS",
                    "",
                    "The following fields appear to contain placeholder text and should be reviewed before PROSPERO submission:",
                    *[f"- {w}" for w in sorted(set(placeholder_warnings))],
                    "",
                ]
            )
        return "\n".join(lines).strip() + "\n"

    def write_prospero_markdown(self, markdown_text: str) -> Path:
        """Write PROSPERO markdown artifact to run output dir."""
        out_dir = self.output_dir
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = out_dir / "doc_prospero_registration.md"
        output_path.write_text(markdown_text, encoding="utf-8")
        return output_path

    def write_markdown(self, workflow_id: str, markdown_text: str) -> Path:
        _ = workflow_id  # Reserved for backward-compatible method signature.
        out_dir = self.output_dir
        out_dir.mkdir(parents=True, exist_ok=True)
        output_path = out_dir / "doc_protocol.md"
        output_path.write_text(markdown_text, encoding="utf-8")
        return output_path
