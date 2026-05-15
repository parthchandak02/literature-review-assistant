from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

from docx import Document

from src.export.docx_exporter import _fix_table_layout, generate_docx


def test_fix_table_layout_bolds_header_row(tmp_path: Path) -> None:
    docx_path = tmp_path / "table.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header A"
    table.cell(0, 1).text = "Header B"
    table.cell(1, 0).text = "Value A"
    table.cell(1, 1).text = "Value B"
    doc.save(docx_path)

    _fix_table_layout(docx_path)

    updated = Document(str(docx_path))
    assert any(run.bold for run in updated.tables[0].rows[0].cells[0].paragraphs[0].runs)


def test_generate_docx_uses_source_resource_path(tmp_path: Path) -> None:
    source = tmp_path / "doc_manuscript.md"
    output = tmp_path / "manuscript.docx"
    source.write_text("# Title\n\n## Introduction\n\nBody\n", encoding="utf-8")

    def _fake_convert_file(source_path: str, to: str, outputfile: str, extra_args: list[str]) -> None:
        assert Path(source_path) == source.resolve()
        assert to == "docx"
        assert "--resource-path" in extra_args
        Document().save(outputfile)

    with patch("src.export.docx_exporter.pypandoc.convert_file", side_effect=_fake_convert_file):
        generated = generate_docx(source, output)

    assert generated == output
    assert output.exists()


def test_generate_docx_rewrites_svg_images_before_conversion(tmp_path: Path) -> None:
    source = tmp_path / "doc_manuscript.md"
    output = tmp_path / "manuscript.docx"
    svg = tmp_path / "fig_concept_taxonomy.svg"
    source.write_text("![Taxonomy](fig_concept_taxonomy.svg)\n", encoding="utf-8")
    svg.write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" width="10" height="10"><rect width="10" height="10"/></svg>',
        encoding="utf-8",
    )

    def _fake_convert_svg(svg_path: Path, png_path: Path) -> None:
        assert svg_path == svg.resolve()
        png_path.write_bytes(b"PNG")

    def _fake_convert_file(source_path: str, to: str, outputfile: str, extra_args: list[str]) -> None:
        assert to == "docx"
        rendered_md = Path(source_path).read_text(encoding="utf-8")
        assert "fig_concept_taxonomy.svg" not in rendered_md
        assert "fig_concept_taxonomy_docx.png" in rendered_md
        Document().save(outputfile)

    with (
        patch("src.export.docx_exporter._convert_svg_to_png", side_effect=_fake_convert_svg),
        patch("src.export.docx_exporter.pypandoc.convert_file", side_effect=_fake_convert_file),
    ):
        generated = generate_docx(source, output)

    assert generated == output
    assert output.exists()
